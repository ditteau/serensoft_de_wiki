#!/usr/bin/env python3
"""
Convert a wiki markdown document to .docx for team distribution.

Handles the subset used in this wiki: headings, paragraphs, bullet and numbered
lists, tables (incl. alignment row), fenced code blocks, blockquotes, horizontal
rules, and inline **bold** / `code` / [text](link).

Usage:
    python tools/md_to_docx.py governance/production-operating-rules.md
    python tools/md_to_docx.py <in.md> --out /path/to/out.docx

The markdown stays the source of truth and lives in git. The .docx is a
distribution artifact -- regenerate it rather than editing it.
"""
import argparse, re, sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY = RGBColor(0x13, 0x40, 0x5A)
MAROON = RGBColor(0x74, 0x00, 0x49)
GREY = RGBColor(0x44, 0x44, 0x44)

INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")


def add_runs(par, text):
    """Render inline bold / code / links into a paragraph."""
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            par.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = par.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
            r.font.color.rgb = MAROON
        elif part.startswith("["):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            r = par.add_run(m.group(1))
            r.font.color.rgb = NAVY
            r.underline = True
        else:
            par.add_run(part)


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def convert(md_path: Path, out_path: Path):
    lines = md_path.read_text().splitlines()
    doc = Document()

    base = doc.styles["Normal"]
    base.font.name = "Verdana"
    base.font.size = Pt(10)

    i = 0
    while i < len(lines):
        line = lines[i]

        # fenced code
        if line.startswith("```"):
            i += 1
            buf = []
            while i < len(lines) and not lines[i].startswith("```"):
                buf.append(lines[i]); i += 1
            i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(10)
            r = p.add_run("\n".join(buf))
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            continue

        # table
        if line.startswith("|") and i + 1 < len(lines) and set(lines[i + 1].replace("|", "").strip()) <= set("-: "):
            header = split_row(line)
            i += 2
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                rows.append(split_row(lines[i])); i += 1
            t = doc.add_table(rows=1, cols=len(header))
            t.style = "Light Grid Accent 1"
            for c, h in enumerate(header):
                cell = t.rows[0].cells[c]
                cell.text = ""
                add_runs(cell.paragraphs[0], h)
                for r_ in cell.paragraphs[0].runs:
                    r_.bold = True
            for row in rows:
                cells = t.add_row().cells
                for c, val in enumerate(row[:len(header)]):
                    cells[c].text = ""
                    add_runs(cells[c].paragraphs[0], val)
            doc.add_paragraph()
            continue

        stripped = line.strip()

        if not stripped:
            i += 1; continue

        if stripped == "---":
            doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1; continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped.lstrip("#").strip()
            h = doc.add_heading(level=min(level, 4))
            add_runs(h, text)
            for r in h.runs:
                r.font.color.rgb = NAVY if level > 1 else MAROON
                r.font.name = "Georgia"
            i += 1; continue

        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            add_runs(p, stripped[2:])
            for r in p.runs:
                r.italic = True
                r.font.color.rgb = GREY
            i += 1; continue

        m = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if m:
            add_runs(doc.add_paragraph(style="List Number"), m.group(2))
            i += 1; continue

        if stripped.startswith(("- ", "* ")):
            add_runs(doc.add_paragraph(style="List Bullet"), stripped[2:])
            i += 1; continue

        add_runs(doc.add_paragraph(), stripped)
        i += 1

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("markdown")
    ap.add_argument("--out", default=None)
    a = ap.parse_args()
    src = Path(a.markdown)
    if not src.exists():
        sys.exit(f"not found: {src}")
    dst = Path(a.out) if a.out else src.with_suffix(".docx")
    convert(src, dst)
    print(f"wrote {dst}  ({dst.stat().st_size:,} bytes)")
